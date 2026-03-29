import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Title Page
title = doc.add_paragraph('BloomFit: AI-Powered Fitness Companion\nAn\nEngineering Project in Community Service\nPhase-2 Individual Report')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.bold = True
    run.font.size = Pt(16)

doc.add_paragraph('\nSubmitted By:').alignment = WD_ALIGN_PARAGRAPH.CENTER
name = doc.add_paragraph('AMAAN SYED\n(23MIP10116)')
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in name.runs:
    run.font.bold = True

doc.add_paragraph('\nSubmitted in partial fulfillment for the award of the degree of\nBachelor of Technology').alignment = WD_ALIGN_PARAGRAPH.CENTER

univ = doc.add_paragraph('\nVIT BHOPAL UNIVERSITY\nBHOPAL, MADHYA PRADESH – 466114')
univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in univ.runs:
    run.font.bold = True

doc.add_paragraph('\nUnder the Supervision of:\nDr. Vijay Kumar Trivedi').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('March-2026\nVIT BHOPAL UNIVERSITY, BHOPAL (M.P.) – 466114').alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# Certificate
cert_title = doc.add_heading('Bonafide Certificate', level=1)
cert_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('This is to certify that this project report titled "BloomFit: AI-Powered Fitness Companion" is the Bonafide work of Amaan Syed (23MIP10116) and he carried out the project work under my supervision as the Guide. This Project Report (Phase-2) is submitted for the Project Phase-2 Examination of Engineering Project in Community Services.')

sigs = doc.add_paragraph('\n\n\nAmaan Syed (23MIP10116)\t\t\tSupervisor (Dr. Vijay Kumar Trivedi)')

doc.add_page_break()

# TOC
toc = doc.add_heading('Table of Contents', level=1)
toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('1. Introduction to BloomFit\n2. Flutter Mobile Application & UI/UX Architecture\n3. React Web Companion & Motion Capture\n4. Cloud Backend & Real-Time Synchronization\n5. Generative AI Integration (Gemini 3 Flash)\n6. Conclusion')

doc.add_page_break()

# Content
doc.add_heading('1. Introduction to BloomFit', level=2)
doc.add_paragraph('The BloomFit project is a multidisciplinary initiative aimed at building a smart, AI-enabled fitness ecosystem capable of autonomous workout generation, real-time form tracking, and cross-device synchronization. As a core member and Lead Developer of the team, I, Amaan Syed, contributed extensively to the frontend architecture, state management, 3D model integration, real-time communication, API interfacing, and the complete construction of the React Web Companion.')
doc.add_paragraph('My role focused primarily on building the core Flutter mobile application using Riverpod, establishing a real-time data bridge via Firebase Firestore, and engineering the React Web Companion for edge-computed motion capture using Google MediaPipe. I also coordinated the prompt engineering required to integrate Google Gemini 3 Flash into our system.')
doc.add_paragraph('This report details my contributions in the domains of software design, system architecture, state management, prototype building, testing, and optimization, supported with observations, results, and learning outcomes.')

doc.add_heading('2. Flutter Mobile Application & UI/UX Architecture', level=2)
doc.add_paragraph('One of my primary responsibilities was to architect and develop the BloomFit mobile application and ensure its reliable performance and seamless user experience.')
doc.add_heading('State Management & Circuit Flow', level=3)
doc.add_paragraph('I began by integrating Riverpod for robust, unidirectional state management. I programmed the active workout session logic to flatten complex, multi-set exercises into a streamlined "Circuit Flow." This allowed the UI to display exactly one set per screen, completely removing visual clutter and cognitive overload for the user. I also implemented a global navigation and dashboard system to track user XP, streaks, and health statistics via Health Connect.')
doc.add_heading('3D Coaching Models', level=3)
doc.add_paragraph('I integrated the model_viewer_plus package to render high-quality .glb 3D character animations natively within the app. I configured the system to dynamically load these heavy 48MB files directly from local bundled assets, ensuring that the animations play instantly and flawlessly without requiring an active internet connection.')

doc.add_heading('3. React Web Companion & Motion Capture', level=2)
doc.add_paragraph('Another significant contribution was the development of mocap-web, a strictly client-side React application acting as a "Smart Mirror."')
doc.add_heading('MediaPipe Integration', level=3)
doc.add_paragraph('I established a computer vision pipeline using Google MediaPipe Pose Landmarker. I programmed the system to continuously process webcam frames entirely on the edge, outputting 33 3D normalized skeletal landmarks at 30+ FPS without ever uploading video data to the cloud, ensuring absolute privacy.')
doc.add_heading('Biomechanical Analysis', level=3)
doc.add_paragraph('I implemented mathematical functions to calculate specific joint angles (e.g., knee depth during a squat). Through rigorous threshold testing, I configured automatic rep counting and form feedback triggering based on these angles, allowing the Web Companion to act as an intelligent, real-time personal trainer.')

doc.add_heading('4. Cloud Backend & Real-Time Synchronization', level=2)
doc.add_paragraph('I designed and deployed the entire serverless Firebase architecture to support the ecosystem.')
doc.add_heading('Database Schema', level=3)
doc.add_paragraph('I structured the Cloud Firestore database, organizing collections for user profiles, historical workout data, and live synchronization sessions.')
doc.add_heading('The Handshake Protocol', level=3)
doc.add_paragraph('I engineered a sub-100ms real-time bridge between the mobile app and the Web Companion. I established a secure protocol where the web app generates a 4-digit code, and the Flutter app updates the corresponding Firestore document. I used onSnapshot listeners and StreamSubscriptions to ensure that as soon as the Web Companion detects a completed rep, the mobile app UI updates instantly.')

doc.add_heading('5. Generative AI Integration (Gemini 3 Flash)', level=2)
doc.add_paragraph('To replace generic workout templates, I developed the PathGenerationService powered by Google Gemini 3 Flash API.')
doc.add_heading('Prompt Engineering & Structured Output', level=3)
doc.add_paragraph('I constructed complex prompts that inject the user\'s primary goal, experience level, and past skipped exercises (e.g., avoiding squats due to injury). I enforced strict JSON schemas on the Gemini response, programming the Flutter app to parse this data into a dynamic, adaptive workout journey that reroutes itself based on the user\'s evolving physical capabilities.')

doc.add_heading('6. Conclusion', level=2)
doc.add_paragraph('Through the development of BloomFit, I gained profound experience in building distributed, real-time ecosystems. The integration of on-device machine learning with generative AI and a highly optimized Flutter frontend resulted in a scalable, privacy-first platform that successfully democratizes access to professional-grade personal training.')

doc.save(r'D:\Programming\02_University\epics\bloomfit\submission-docs\EPICS_Individual_Report_Amaan_Syed.docx')
